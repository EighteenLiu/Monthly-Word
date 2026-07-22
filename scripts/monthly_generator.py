"""
Generate the monthly sealed cleaning station inspection bulletin.

The script first checks whether raw .doc daily reports have already been
converted to .docx. Existing converted files are reused when they are newer
than the corresponding raw file; only missing or stale files are converted.
Then it summarizes all converted daily reports for the requested month and
creates the monthly .docx bulletin.
"""

from __future__ import annotations

import argparse
import io
import re
import shutil
import subprocess
import sys
import uuid
from collections import defaultdict
from dataclasses import dataclass, field
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from jinja2 import Environment
from PIL import Image


def get_project_root() -> Path:
    if getattr(sys, "frozen", False):
        return Path(sys.executable).resolve().parent
    return Path(__file__).resolve().parents[1]


PROJECT_ROOT = get_project_root()
DEFAULT_RAW_INPUT_DIR = PROJECT_ROOT / "01_原始日报"
DEFAULT_CONVERTED_DIR = PROJECT_ROOT / "02_转换后日报"
DEFAULT_TEMPLATE_DIR = PROJECT_ROOT / "03_月报模板"
DEFAULT_OUTPUT_DIR = PROJECT_ROOT / "04_输出月报"
DEFAULT_STATION_TYPE = "清洁站"
DEFAULT_YEAR = 2026
TEMP_WORK_ROOT_NAME = "_monthly_work"
A4_WIDTH = Cm(21)
A4_HEIGHT = Cm(29.7)

DAILY_NAME_RE = re.compile(r"(?P<month>\d{1,2})月(?P<day>\d{1,2})日")
DAILY_CODE_RE = re.compile(r"(?P<year>20\d{2})(?P<month>\d{2})(?P<day>\d{2})")

CLEANING_ISSUE_CATEGORIES: list[tuple[str, tuple[str, ...]]] = [
    ("无称重系统或称重系统损坏", ("无称重系统", "称重系统损坏", "无称重", "称重设备损坏", "称重屏幕破损")),
    ("小型收集车混装混运", ("小型收集车混装混运", "混装混运")),
    ("箱体内垃圾混投", ("箱体内垃圾混投", "垃圾混投", "混投")),
    ("未开门运行", ("未开门运行", "未开门", "不运行", "未运行")),
    ("拒收单不准确", ("拒收单不准确", "拒收单")),
]

TRANSFER_ISSUE_CATEGORIES: list[tuple[str, tuple[str, ...]]] = [
    ("清运不及时、可回收物大量积存", ("清运不及时", "大量积存", "可回收物大量积存")),
    ("周边环境脏乱", ("周边环境脏乱", "环境脏乱")),
    ("无可回收价格表或价格表损坏", ("无可回收价格表", "无价格表", "价格表损坏", "价格表破损", "无可回收物价格表", "价目表", "价目价目表")),
    ("无备案公示", ("无备案公示", "无备案信息", "备案信息表", "备案公示")),
    ("消防水源不合格", ("消防水源不合格",)),
    ("无消防安全水源", ("无消防安全水源", "无消防水源")),
    ("无营业执照", ("无营业执照",)),
    ("无安全风险公告", ("无安全风险公告", "无公示牌", "安全风险公告")),
    ("称重系统损坏", ("称重系统损坏", "称重屏幕破损", "称重设备损坏", "计量称不能使用")),
    ("灭火器过期", ("灭火器过期",)),
    ("灭火器不合格", ("灭火器不合格",)),
    ("未按规定区域存放物品", ("未按规定区域存放", "未按规定存放", "暂存区未按规定存放", "堆放混乱")),
    ("无七禁收八不准承诺书", ("无七禁收八不准承诺书", "七禁收八不准")),
    ("配电箱处堆放杂物", ("配电箱处堆放杂物", "配电箱堆放杂物")),
    ("安全员未按时上岗", ("安全员未按时上岗",)),
    ("安全员无明显身份标识", ("安全员无明显身份标识", "全员无明显身份标识", "无明显身份标识")),
    ("未按时开门运行", ("未按时开门运行", "未开门运行", "未开门", "不运行", "未运行")),
    ("无企安安", ("无企安安", "企安安无法正常登录")),
    ("无灭蝇措施", ("无灭蝇措施",)),
]

STATION_PROFILES = {
    "清洁站": {
        "title": "西城区密闭式清洁站存在问题检查通报",
        "output_name": "西城区{year}年{month}月密闭式清洁站检查通报.docx",
        "subject": "密闭式清洁站",
        "attachment_subject": "密闭式清洁站",
        "detail_heading": "二、各街道情况",
        "categories": CLEANING_ISSUE_CATEGORIES,
        "unopened_category": "未开门运行",
        "non_problem_status_keywords": ("已关闭", "已关门", "已停业", "停业", "闭店", "歇业", "暂停营业"),
    },
    "中转站": {
        "title": "西城区可回收物中转站检查情况通报",
        "output_name": "西城区{year}年{month}月中转站检查通报.docx",
        "subject": "可回收物中转站",
        "attachment_subject": "可回收物中转站",
        "detail_heading": "二、各街道案例",
        "categories": TRANSFER_ISSUE_CATEGORIES,
        "unopened_category": "未按时开门运行",
        "non_problem_status_keywords": ("升级改造", "施工停运", "停运"),
    },
}


class ConversionError(RuntimeError):
    """Raised when a daily report cannot be converted."""


@dataclass(frozen=True)
class ImageSpec:
    blob: bytes
    width_pt: float | None
    height_pt: float | None


@dataclass(frozen=True)
class ImageRow:
    images: tuple[ImageSpec, ...]
    centered: bool = True


@dataclass(frozen=True)
class DailyRecord:
    source: Path
    month: int
    day: int
    street: str
    station: str
    issue_text: str
    issue_counts: dict[str, int]
    image_rows: tuple[ImageRow, ...] = ()

    @property
    def date_label(self) -> str:
        return f"{self.month}月{self.day}日"

    @property
    def has_problem(self) -> bool:
        return normalize_issue_text(self.issue_text) != "无问题。"


@dataclass
class MergedRecord:
    station: str
    issue_texts: list[str]
    image_rows: list[ImageRow]


@dataclass
class ConversionSummary:
    converted: int = 0
    skipped: int = 0
    skipped_out_of_period: int = 0
    failed: list[tuple[Path, str]] = field(default_factory=list)


def iter_doc_files(input_dir: Path) -> list[Path]:
    if not input_dir.exists():
        return []
    return sorted(
        path
        for path in input_dir.rglob("*.doc")
        if path.is_file() and path.suffix.lower() == ".doc" and not path.name.startswith("~$")
    )


def iter_docx_files(input_dir: Path) -> list[Path]:
    if not input_dir.exists():
        return []
    return sorted(
        path
        for path in input_dir.rglob("*.docx")
        if path.is_file() and path.suffix.lower() == ".docx" and not path.name.startswith("~$")
    )


def find_soffice() -> str | None:
    for executable in ("soffice", "libreoffice"):
        found = shutil.which(executable)
        if found:
            return found
    return None


def convert_with_word(source: Path, target: Path) -> None:
    try:
        import pythoncom
        import win32com.client
    except ImportError as exc:
        raise ConversionError("pywin32 is not installed") from exc

    pythoncom.CoInitialize()
    word = None
    document = None
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        document = word.Documents.Open(str(source.resolve()))
        target.parent.mkdir(parents=True, exist_ok=True)
        document.SaveAs2(str(target.resolve()), FileFormat=16)
    except Exception as exc:
        raise ConversionError(str(exc)) from exc
    finally:
        if document is not None:
            document.Close(False)
        if word is not None:
            word.Quit()
        pythoncom.CoUninitialize()


def convert_with_libreoffice(source: Path, target: Path) -> None:
    soffice = find_soffice()
    if not soffice:
        raise ConversionError("LibreOffice/soffice was not found in PATH")

    target.parent.mkdir(parents=True, exist_ok=True)
    result = subprocess.run(
        [
            soffice,
            "--headless",
            "--convert-to",
            "docx",
            "--outdir",
            str(target.parent.resolve()),
            str(source.resolve()),
        ],
        capture_output=True,
        text=True,
        check=False,
    )
    generated = target.parent / f"{source.stem}.docx"
    if result.returncode != 0 or not generated.exists():
        message = (result.stderr or result.stdout or "LibreOffice conversion failed").strip()
        raise ConversionError(message)
    if generated.resolve() != target.resolve():
        generated.replace(target)


def convert_one(source: Path, target: Path, engine: str) -> bool:
    if target.exists() and target.stat().st_mtime >= source.stat().st_mtime:
        return False

    engines = ["word", "libreoffice"] if engine == "auto" else [engine]
    errors: list[str] = []
    for selected in engines:
        try:
            if selected == "word":
                convert_with_word(source, target)
            elif selected == "libreoffice":
                convert_with_libreoffice(source, target)
            else:
                raise ConversionError(f"Unknown conversion engine: {selected}")
            return True
        except ConversionError as exc:
            errors.append(f"{selected}: {exc}")

    raise ConversionError("; ".join(errors))


def ensure_converted(
    raw_input_dir: Path,
    converted_dir: Path,
    station_type: str = DEFAULT_STATION_TYPE,
    engine: str = "auto",
    year: int | None = None,
    month: int | None = None,
) -> ConversionSummary:
    station_type = normalize_station_type(station_type)
    target_root = converted_dir / station_type
    summary = ConversionSummary()
    period: tuple[date, date] | None = report_period(year, month) if year is not None and month is not None else None

    for source in iter_doc_files(raw_input_dir):
        if not path_matches_station_type(source, station_type):
            continue
        if period is not None:
            parsed_date = daily_date_from_name(source)
            if parsed_date is None:
                summary.skipped_out_of_period += 1
                continue
            record_date = record_date_for_period(parsed_date[0], parsed_date[1], period[0], period[1])
            if not period[0] <= record_date <= period[1]:
                summary.skipped_out_of_period += 1
                continue
        relative = source.relative_to(raw_input_dir)
        target = (target_root / relative).with_suffix(".docx")
        try:
            was_converted = convert_one(source, target, engine)
        except ConversionError as exc:
            summary.failed.append((source, str(exc)))
            continue

        if was_converted:
            summary.converted += 1
            print(f"[转换] {source.name} -> {target.name}")
        else:
            summary.skipped += 1

    if summary.skipped:
        print(f"[跳过] {summary.skipped} 个日报已完成 doc -> docx 转换")
    if summary.skipped_out_of_period:
        print(f"[跳过] {summary.skipped_out_of_period} 个日报不在本次报告期内，未转换")
    if summary.converted:
        print(f"[完成] 本次新转换 {summary.converted} 个日报")
    if not summary.converted and not summary.skipped and not summary.skipped_out_of_period:
        matching_docx = [path for path in iter_docx_files(raw_input_dir) if path_matches_station_type(path, station_type)]
        if matching_docx:
            print(f"[提示] 未发现原始 .doc 日报，将直接汇总 {len(matching_docx)} 个已生成 .docx 日报")
        else:
            print(f"[提示] 未发现 {station_type} 原始 .doc 或已生成 .docx 日报：{raw_input_dir}")
    return summary


def is_on_c_drive(path: Path) -> bool:
    drive = path.resolve().drive
    return drive.lower() == "c:"


def create_monthly_work_dir(*candidate_roots: Path) -> Path:
    for root in candidate_roots:
        if not root:
            continue
        resolved_root = root.resolve()
        if is_on_c_drive(resolved_root):
            continue
        work_dir = resolved_root / TEMP_WORK_ROOT_NAME / uuid.uuid4().hex
        work_dir.mkdir(parents=True, exist_ok=False)
        return work_dir
    raise RuntimeError("无法创建临时目录：候选路径都位于 C 盘，请选择非 C 盘输出目录。")


def cleanup_work_dir(work_dir: Path | None) -> None:
    if work_dir is not None:
        shutil.rmtree(work_dir, ignore_errors=True)


def parse_month(value: str | int | None) -> int:
    if value is None:
        raise ValueError("请通过 --month 指定月份，例如 4 或 2026-04")
    text = str(value).strip()
    match = re.search(r"(?:(?:\d{4})[-年])?(?P<month>\d{1,2})(?:月)?$", text)
    if not match:
        raise ValueError(f"无法识别月份：{value}")
    month = int(match.group("month"))
    if not 1 <= month <= 12:
        raise ValueError(f"月份超出范围：{value}")
    return month


def parse_year(value: int | None, month_value: str | int | None) -> int:
    if value:
        return value
    if month_value:
        match = re.search(r"(?P<year>\d{4})[-年]", str(month_value))
        if match:
            return int(match.group("year"))
    return DEFAULT_YEAR


def previous_month(year: int, month: int) -> tuple[int, int]:
    if month == 1:
        return year - 1, 12
    return year, month - 1


def report_period(year: int, month: int) -> tuple[date, date]:
    start_year, start_month = previous_month(year, month)
    return date(start_year, start_month, 20), date(year, month, 19)


def record_date_for_period(record_month: int, record_day: int, start: date, end: date) -> date:
    if record_month == start.month:
        return date(start.year, record_month, record_day)
    if record_month == end.month:
        return date(end.year, record_month, record_day)
    return date(end.year, record_month, record_day)


def format_date_cn(value: date) -> str:
    return f"{value.year}年{value.month}月{value.day}日"


def daily_date_from_name(path: Path) -> tuple[int, int] | None:
    match = DAILY_NAME_RE.search(path.stem)
    if match:
        return int(match.group("month")), int(match.group("day"))
    match = DAILY_CODE_RE.search(path.stem)
    if match:
        return int(match.group("month")), int(match.group("day"))
    return None


def normalize_issue_text(text: str) -> str:
    text = re.sub(r"\s+", "", text or "")
    text = re.sub(r"^存在的问题是[:：]?", "", text)
    if not text or text in {"无问题", "未发现问题"}:
        return "无问题。"
    return text if text.endswith(("。", "！", "？")) else f"{text}。"


def strip_heading_number_prefix(text: str) -> str:
    text = (text or "").strip()
    number_pattern = re.compile(
        r"^\s*(?:"
        r"[（(][一二三四五六七八九十百千万零〇\d]+[）)]"
        r"|[一二三四五六七八九十百千万零〇\d]+[、.．]"
        r")\s*"
    )
    while True:
        stripped = number_pattern.sub("", text, count=1).strip()
        if stripped == text:
            return stripped
        text = stripped


def normalize_street_name(value: str) -> str:
    value = strip_heading_number_prefix(value).strip()
    # 广外 = 广安门外，广内 = 广安门内，统一用简称避免同一街道出现重复标题
    value = value.replace("广安门外", "广外")
    value = value.replace("广安门内", "广内")
    return value


def split_street_station(title: str) -> tuple[str, str]:
    title = strip_heading_number_prefix(title)
    marker = "街道"
    idx = title.find(marker)
    if idx == -1:
        return "未识别街道", title.strip()
    street = normalize_street_name(title[: idx + len(marker)])
    station = title[idx + len(marker) :].strip()
    return street, station or title.strip()


def get_station_profile(station_type: str) -> dict:
    aliases = {
        "1": "清洁站",
        "cleaning": "清洁站",
        "clean": "清洁站",
        "sealed": "清洁站",
        "2": "中转站",
        "transfer": "中转站",
        "recycle": "中转站",
    }
    station_type = aliases.get(station_type, station_type)
    if station_type not in STATION_PROFILES:
        supported = "、".join(STATION_PROFILES)
        raise ValueError(f"不支持的日报类型：{station_type}。目前支持：{supported}")
    return STATION_PROFILES[station_type]


def normalize_station_type(station_type: str) -> str:
    aliases = {
        "1": "清洁站",
        "cleaning": "清洁站",
        "clean": "清洁站",
        "sealed": "清洁站",
        "2": "中转站",
        "transfer": "中转站",
        "recycle": "中转站",
    }
    return aliases.get(station_type, station_type)


def infer_station_type_from_path(path: Path) -> str | None:
    filename_type = infer_station_type_from_filename(path)
    if filename_type:
        return filename_type
    for part in reversed(path.parts[:-1]):
        if part in {"清洁站", "密闭式清洁站"}:
            return "清洁站"
        if part in {"中转站", "可回收物中转站"}:
            return "中转站"
    return None


def infer_station_type_from_filename(path: Path) -> str | None:
    filename = path.name
    is_clean = "密闭式清洁站" in filename or "清洁站" in filename
    is_transfer = "可回收物中转站" in filename or "中转站" in filename
    if is_clean == is_transfer:
        return None
    return "清洁站" if is_clean else "中转站"


def path_matches_station_type(path: Path, station_type: str) -> bool:
    return infer_station_type_from_path(path) == normalize_station_type(station_type)


def count_issue_categories(issue_text: str, categories: list[tuple[str, tuple[str, ...]]]) -> dict[str, int]:
    normalized = normalize_issue_text(issue_text)
    counts = {category: 0 for category, _ in categories}
    if normalized == "无问题。":
        return counts
    for item in split_issue_items(normalized):
        category = best_issue_category(item, categories)
        if category:
            counts[category] = 1
    return counts


def best_issue_category(item: str, categories: list[tuple[str, tuple[str, ...]]]) -> str | None:
    normalized_item = normalize_issue_match_text(item)
    best_category: str | None = None
    best_length = -1
    for category, keywords in categories:
        for keyword in (category, *keywords):
            for candidate in issue_match_candidates(keyword):
                if candidate in normalized_item and len(candidate) > best_length:
                    best_category = category
                    best_length = len(candidate)
    return best_category


def normalize_issue_match_text(text: str) -> str:
    return re.sub(r"\s+", "", text or "").replace("台帐", "台账").replace("帐", "账")


def issue_match_candidates(text: str) -> tuple[str, ...]:
    normalized = normalize_issue_match_text(text).strip("，,。；;：:")
    if not normalized:
        return ()

    candidates = {normalized}
    if normalized.startswith("无") and len(normalized) > 1:
        candidates.add(normalized[1:])
    for candidate in tuple(candidates):
        if "精细化" in candidate:
            candidates.add(candidate.replace("精细化", "细化"))

    return tuple(sorted(candidates, key=len, reverse=True))


def previous_nonempty_paragraph_text(paragraphs, index: int) -> str:
    for cursor in range(index - 1, -1, -1):
        text = paragraphs[cursor].text.strip()
        if text:
            return text
    return ""


def is_overall_heading(text: str) -> bool:
    return text in {"整体情况", "1.整体情况", "一、整体情况"} or text.endswith("整体情况")


def is_detail_heading(text: str) -> bool:
    return text in {"具体情况", "2.具体情况", "二、具体情况"} or text.endswith("具体情况")


def next_overall_index(paragraphs, index: int) -> int:
    for cursor in range(index + 1, len(paragraphs)):
        if is_overall_heading(paragraphs[cursor].text.strip()):
            return cursor
    return len(paragraphs)


def parse_points_from_style(style: str) -> tuple[float | None, float | None]:
    width_match = re.search(r"width:([0-9.]+)pt", style or "")
    height_match = re.search(r"height:([0-9.]+)pt", style or "")
    width = float(width_match.group(1)) if width_match else None
    height = float(height_match.group(1)) if height_match else None
    return width, height


def image_row_from_paragraph(document: Document, paragraph) -> ImageRow | None:
    images: list[ImageSpec] = []
    seen_rids: set[str] = set()
    shapes = paragraph._p.xpath('.//*[local-name()="shape"]')
    for shape in shapes:
        image_data = shape.xpath('.//*[local-name()="imagedata"]')
        if not image_data:
            continue
        rid = image_data[0].get(qn("r:id"))
        if not rid or rid in seen_rids:
            continue
        seen_rids.add(rid)
        style = shape.get("style", "")
        related_part = document.part.related_parts.get(rid)
        if related_part is None:
            continue
        width, height = parse_points_from_style(style)
        images.append(ImageSpec(blob=related_part.blob, width_pt=width, height_pt=height))
    blips = paragraph._p.xpath('.//*[local-name()="blip"]')
    for blip in blips:
        rid = blip.get(qn("r:embed")) or blip.get(qn("r:link"))
        if not rid or rid in seen_rids:
            continue
        seen_rids.add(rid)
        related_part = document.part.related_parts.get(rid)
        if related_part is None:
            continue
        images.append(ImageSpec(blob=related_part.blob, width_pt=None, height_pt=None))
    if not images:
        return None
    return ImageRow(images=tuple(images), centered=paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER)


def detail_heading_matches_issue(heading: str, issue_text: str) -> bool:
    issue = normalize_issue_text(issue_text)
    rules = [
        (("混投", "垃圾", "箱体"), ("内部", "箱体", "垃圾")),
        (("运输", "喷涂", "车辆"), ("运输", "车辆")),
        (("拒收单",), ("拒收", "证件", "文件", "内部")),
        (("称重",), ("称重", "电箱", "公示", "备案", "证件", "文件", "内部")),
        (("价格表", "价目表", "价目"), ("价格", "价目", "公示", "公告", "营业执照")),
        (("备案", "公示牌", "安全风险公告", "七禁收八不准", "企安安", "营业执照"), ("备案", "公告", "公示", "营业执照", "承诺书", "证件", "负责人", "安全制度")),
        (("安全员", "身份标识"), ("负责人", "证件", "安全员", "门牌", "正门")),
        (("周边环境", "环境脏乱"), ("周边环境", "内部环境", "内部及环境")),
        (("消防水源",), ("消防", "水源", "内部")),
        (("配电箱", "电源箱"), ("配电箱", "电源箱", "内部")),
        (("堆放混乱", "未按规定", "暂存区"), ("内部环境", "内部", "暂存")),
        (("灭火器",), ("灭火器",)),
        (("门牌", "正门"), ("正门", "门牌")),
    ]
    for issue_keywords, heading_keywords in rules:
        if any(keyword in issue for keyword in issue_keywords):
            return any(keyword in heading for keyword in heading_keywords)
    return False


def extract_problem_image_rows(document: Document, detail_start: int, detail_end: int, issue_text: str) -> tuple[ImageRow, ...]:
    sections: list[tuple[str, list[ImageRow]]] = []
    current_heading = ""
    current_rows: list[ImageRow] = []

    for paragraph in document.paragraphs[detail_start:detail_end]:
        text = paragraph.text.strip()
        row = image_row_from_paragraph(document, paragraph)
        if text:
            if current_heading or current_rows:
                sections.append((current_heading, current_rows))
            current_heading = text
            current_rows = []
        elif row is not None:
            current_rows.append(row)

    if current_heading or current_rows:
        sections.append((current_heading, current_rows))

    all_rows = [row for _, rows in sections for row in rows]
    issue_items = split_issue_items(issue_text)
    if not issue_items:
        return tuple(all_rows)

    matched_rows: list[ImageRow] = []
    seen_signatures: set[tuple[tuple[int, bytes], ...]] = set()
    for item in issue_items:
        item_rows = [
            row
            for heading, rows in sections
            if detail_heading_matches_issue(heading, item)
            for row in rows
        ]
        if not item_rows:
            return tuple(all_rows)
        for row in item_rows:
            signature = image_row_signature(row)
            if signature in seen_signatures:
                continue
            matched_rows.append(row)
            seen_signatures.add(signature)

    return tuple(matched_rows)


def parse_daily_report(path: Path, categories: list[tuple[str, tuple[str, ...]]]) -> list[DailyRecord]:
    date = daily_date_from_name(path)
    if date is None:
        return []
    month, day = date
    document = Document(str(path))
    paragraphs = [p.text.strip() for p in document.paragraphs]

    records: list[DailyRecord] = []
    for index, text in enumerate(paragraphs):
        if not is_overall_heading(text) or index == 0:
            continue

        title = previous_nonempty_paragraph_text(document.paragraphs, index)
        issue_parts: list[str] = []
        cursor = index + 1
        while cursor < len(paragraphs) and not is_detail_heading(paragraphs[cursor]):
            if paragraphs[cursor]:
                issue_parts.append(paragraphs[cursor])
            cursor += 1

        street, station = split_street_station(title)
        issue_text = normalize_issue_text("".join(issue_parts))
        detail_start = cursor + 1 if cursor < len(paragraphs) else cursor
        detail_end = max(detail_start, next_overall_index(document.paragraphs, index))
        records.append(
            DailyRecord(
                source=path,
                month=month,
                day=day,
                street=street,
                station=station,
                issue_text=issue_text,
                issue_counts=count_issue_categories(issue_text, categories),
                image_rows=extract_problem_image_rows(document, detail_start, detail_end, issue_text),
            )
        )
    return records


def collect_records(source_dirs: Path | list[Path], station_type: str, year: int, month: int) -> list[DailyRecord]:
    station_type = normalize_station_type(station_type)
    if isinstance(source_dirs, Path):
        source_dirs = [source_dirs]
    categories = get_station_profile(station_type)["categories"]
    start, end = report_period(year, month)
    records: list[DailyRecord] = []
    seen: set[Path] = set()
    paths: list[Path] = []
    skipped_other_type = 0
    for source_dir in source_dirs:
        for path in iter_docx_files(source_dir):
            resolved = path.resolve()
            if resolved in seen:
                continue
            seen.add(resolved)
            paths.append(path)
    for path in sorted(paths):
        filename_type = infer_station_type_from_filename(path)
        if filename_type and filename_type != station_type:
            skipped_other_type += 1
            continue
        if filename_type is None and not path_matches_station_type(path, station_type):
            continue
        parsed_date = daily_date_from_name(path)
        if parsed_date is None:
            continue
        record_date = record_date_for_period(parsed_date[0], parsed_date[1], start, end)
        if not start <= record_date <= end:
            continue
        records.extend(parse_daily_report(path, categories))
    if skipped_other_type:
        print(f"[跳过] {skipped_other_type} 个文件名类型与当前月报类型不一致的日报")
    return sorted(
        records,
        key=lambda item: (record_date_for_period(item.month, item.day, start, end), item.street, item.station),
    )


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = borders.find(qn(f"w:{border_name}"))
        if border is None:
            border = OxmlElement(f"w:{border_name}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_table_row_min_height(row, height_cm: float = 0.8) -> None:
    row.height = Cm(height_cm)
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "仿宋_GB2312"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")
    run.font.size = Pt(10)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def remove_empty_paragraph(paragraph) -> None:
    paragraph._element.getparent().remove(paragraph._element)
    paragraph._p = paragraph._element = None


def cleanup_empty_paragraphs(document: Document) -> None:
    for paragraph in list(document.paragraphs):
        xml = paragraph._p.xml
        if paragraph.text.strip():
            continue
        if "a:blip" in xml or "v:imagedata" in xml or "w:sectPr" in xml:
            continue
        remove_empty_paragraph(paragraph)


def clear_document_body(document: Document) -> None:
    body = document._body._element
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def create_document_from_template(template_path: Path | None) -> Document:
    if template_path and template_path.exists():
        document = Document(str(template_path))
        clear_document_body(document)
        return document
    return Document()


def configure_portrait_a4_section(section) -> None:
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = A4_WIDTH
    section.page_height = A4_HEIGHT
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)


def configure_landscape_a4_section(section) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = A4_HEIGHT
    section.page_height = A4_WIDTH
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)


def add_title(document: Document, text: str) -> None:
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(18)
    run = title.add_run(text)
    run.bold = True
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(22)


def read_docx_template_lines(template_path: Path | None, profile: dict) -> list[str]:
    if not template_path or not template_path.exists():
        return []
    source = Document(str(template_path))
    lines = [paragraph.text.strip() for paragraph in source.paragraphs if paragraph.text.strip()]
    if not lines:
        return []

    parsed: list[str] = []
    saw_summary = False
    saw_detail = False
    for index, line in enumerate(lines):
        if index == 0:
            parsed.append("{{ title }}")
            continue
        if "总体情况" in line and not saw_summary:
            parsed.extend(["", "{{ summary_heading }}", "{{ summary }}"])
            saw_summary = True
            continue
        if ("各街道" in line or "街道案例" in line) and not saw_detail:
            parsed.extend(["", "{{ detail_heading }}", "{% if streets %}"])
            parsed.extend(default_street_template_lines())
            parsed.extend(["{% else %}", "{{ no_detail_text }}", "{% endif %}"])
            saw_detail = True
            break
        if "附件" in line:
            break

    if saw_summary and saw_detail:
        return parsed
    return []


def default_street_template_lines() -> list[str]:
    return [
        "{% for street in streets %}",
        "{{ street.heading }}",
        "{% for station in street.stations %}",
        "[[STATION:{{ station.key }}]]",
        "[[IMAGES:{{ station.image_key }}]]",
        "{% endfor %}",
        "{% endfor %}",
    ]


def default_jinja_template_content(profile: dict) -> str:
    lines = [
        "{{ title }}",
        "",
        "{{ summary_heading }}",
        "{{ summary }}",
        "",
        "{{ detail_heading }}",
        "{% if streets %}",
        *default_street_template_lines(),
        "{% else %}",
        "{{ no_detail_text }}",
        "{% endif %}",
        "",
    ]
    return "\n".join(lines)


def export_jinja_text_template(
    template_path: Path | None,
    station_type: str,
    profile: dict,
    work_dir: Path | None = None,
) -> Path | None:
    if not template_path:
        return None
    target_dir = work_dir or create_monthly_work_dir(DEFAULT_OUTPUT_DIR, PROJECT_ROOT)
    target_dir.mkdir(parents=True, exist_ok=True)
    station_name = normalize_station_type(station_type)
    target = target_dir / f"{station_name}_auto_template.jinja2"
    parsed_lines = read_docx_template_lines(template_path, profile)
    content = "\n".join(parsed_lines) if parsed_lines else default_jinja_template_content(profile)
    target.write_text(content, encoding="utf-8")
    return target


def build_template_context(
    records: list[DailyRecord],
    year: int,
    month: int,
    station_type: str,
    profile: dict,
    categories: list[tuple[str, tuple[str, ...]]],
) -> tuple[dict, dict[str, tuple[ImageRow, ...]], dict[str, dict[str, int]]]:
    body_records = build_body_records(records, station_type, categories)
    by_street, counts = summarize_records(body_records, categories)

    image_map: dict[str, tuple[ImageRow, ...]] = {}
    station_map: dict[str, dict] = {}
    streets: list[dict] = []
    for street_index, (street, street_records) in enumerate(by_street.items(), start=1):
        use_station_numbers = len(street_records) > 1
        stations: list[dict] = []
        for station_index, record in enumerate(street_records, start=1):
            issue_text, has_problem = format_station_issue_text(record, profile)
            station_label = format_station_label(record, station_index, use_station_numbers, station_type)
            image_key = f"s{street_index}_p{station_index}"
            station_key = image_key
            image_map[image_key] = (
                tuple(record.image_rows) if has_problem else limit_image_rows(tuple(record.image_rows), 3)
            )
            station_data = {
                "key": station_key,
                "name": record.station,
                "label": station_label,
                "issue_text": issue_text,
                "text": f"{station_label}：{issue_text}",
                "image_key": image_key,
                "has_problem": has_problem,
            }
            station_map[station_key] = station_data
            stations.append(station_data)
        streets.append(
            {
                "name": street,
                "index": street_index,
                "heading": f"（{chinese_section_number(street_index)}）{street}",
                "stations": stations,
            }
        )

    start, end = report_period(year, month)
    context = {
        "title": profile["title"],
        "summary_heading": "一、总体情况",
        "summary": build_summary_sentence(body_records, year, month, profile, len(by_street)),
        "detail_heading": profile["detail_heading"],
        "no_detail_text": "本月未发现需列入街道情况的站点。",
        "streets": streets,
        "station_map": station_map,
        "start_date": format_date_cn(start),
        "end_date": format_date_cn(end),
        "year": year,
        "month": month,
        "subject": profile["subject"],
        "attachment_subject": profile["attachment_subject"],
    }
    return context, image_map, counts


def render_jinja_body(document: Document, jinja_path: Path | None, context: dict, image_map: dict[str, tuple[ImageRow, ...]]) -> None:
    template_text = jinja_path.read_text(encoding="utf-8") if jinja_path and jinja_path.exists() else default_jinja_template_content({})
    rendered = Environment(autoescape=False, trim_blocks=True, lstrip_blocks=True).from_string(template_text).render(**context)
    for raw_line in rendered.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        station_match = re.fullmatch(r"\[\[STATION:(?P<key>[A-Za-z0-9_:-]+)\]\]", line)
        if station_match:
            station = context["station_map"].get(station_match.group("key"))
            if station:
                add_station_paragraph(document, station["label"], station["issue_text"], station["has_problem"])
            continue
        image_match = re.fullmatch(r"\[\[IMAGES:(?P<key>[A-Za-z0-9_:-]+)\]\]", line)
        if image_match:
            add_image_rows(document, image_map.get(image_match.group("key"), ()))
            continue
        if line == context["title"]:
            add_title(document, line)
        elif line in (context["summary_heading"], context["detail_heading"]) or re.match(r"^[一二三四五六七八九十]+、", line):
            add_heading(document, line, level=1)
        elif re.match(r"^（[一二三四五六七八九十]+）", line):
            add_heading(document, line, level=2)
        else:
            add_paragraph(document, line)


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "仿宋"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
    normal.font.size = Pt(16)

    title = get_or_create_paragraph_style(document, ("Title", "标题"), "Title")
    if title:
        title.font.name = "黑体"
        title._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        title.font.size = Pt(18)
        title.font.bold = True

    heading1 = get_or_create_paragraph_style(document, ("Heading 1", "标题 1", "标题1", "Heading1"), "Heading 1")
    if heading1:
        heading1.font.name = "黑体"
        heading1._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        heading1.font.size = Pt(16)
        heading1.font.bold = True
        heading1.font.color.rgb = RGBColor(0, 0, 0)
        set_style_outline_level(heading1, 0)

    heading2 = get_or_create_paragraph_style(document, ("Heading 2", "标题 2", "标题2", "Heading2"), "Heading 2")
    if heading2:
        heading2.font.name = "楷体_GB2312"
        heading2._element.rPr.rFonts.set(qn("w:eastAsia"), "楷体_GB2312")
        heading2.font.size = Pt(16)
        heading2.font.bold = True
        heading2.font.color.rgb = RGBColor(0, 0, 0)
        set_style_outline_level(heading2, 1)

    heading3 = get_or_create_paragraph_style(document, ("Heading 3", "标题 3", "标题3", "Heading3"), "Heading 3")
    if heading3:
        heading3.font.name = "仿宋"
        heading3._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
        heading3.font.size = Pt(16)
        heading3.font.bold = False
        heading3.font.color.rgb = RGBColor(0, 0, 0)
        set_style_outline_level(heading3, 2)


def get_first_style(document: Document, names: tuple[str, ...]):
    for name in names:
        try:
            return document.styles[name]
        except KeyError:
            continue
    for style in document.styles:
        if style.style_id in names:
            return style
    return None


def get_or_create_paragraph_style(document: Document, names: tuple[str, ...], fallback_name: str):
    style = get_first_style(document, names)
    if style:
        return style
    return document.styles.add_style(fallback_name, WD_STYLE_TYPE.PARAGRAPH)


def set_style_outline_level(style, level: int) -> None:
    p_pr = style._element.get_or_add_pPr()
    remove_numbering_from_ppr(p_pr)
    outline = p_pr.find(qn("w:outlineLvl"))
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        p_pr.append(outline)
    outline.set(qn("w:val"), str(level))
    q_format = style._element.find(qn("w:qFormat"))
    if q_format is None:
        style._element.append(OxmlElement("w:qFormat"))


def get_heading_style_name(document: Document, level: int) -> str | None:
    style = get_or_create_paragraph_style(
        document,
        (f"Heading {level}", f"标题 {level}", f"标题{level}", f"Heading{level}"),
        f"Heading {level}",
    )
    return style.name if style else None


def remove_numbering_from_ppr(p_pr) -> None:
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is not None:
        p_pr.remove(num_pr)


def remove_paragraph_numbering(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    remove_numbering_from_ppr(p_pr)


def add_paragraph(document: Document, text: str, *, bold: bool = False, align=None) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = None if bold else Cm(0.74)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(0)
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "黑体" if bold else "仿宋"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体" if bold else "仿宋")
    run.font.size = Pt(16)


def add_station_paragraph(document: Document, label: str, issue_text: str, has_problem: bool) -> None:
    paragraph = document.add_paragraph(style=get_heading_style_name(document, 3))
    remove_paragraph_numbering(paragraph)
    paragraph.paragraph_format.first_line_indent = Cm(0.74)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(0)

    label_run = paragraph.add_run(f"{label}：")
    label_run.font.name = "仿宋"
    label_run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
    label_run.font.size = Pt(16)

    issue_run = paragraph.add_run(issue_text)
    issue_run.font.name = "仿宋"
    issue_run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
    issue_run.font.size = Pt(16)
    if has_problem:
        issue_run.font.color.rgb = RGBColor(255, 0, 0)


def add_image_rows(document: Document, image_rows: tuple[ImageRow, ...]) -> None:
    images = [image for image_row in image_rows for image in image_row.images]
    image_width = Cm(4.55)
    for index in range(0, len(images), 3):
        row_images = images[index : index + 3]
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        for image in row_images:
            run = paragraph.add_run()
            image_stream = io.BytesIO()
            with Image.open(io.BytesIO(image.blob)) as pil_image:
                pil_image.save(image_stream, format="PNG")
            image_stream.seek(0)
            run.add_picture(image_stream, width=image_width)


def add_heading(document: Document, text: str, level: int) -> None:
    paragraph = document.add_paragraph(style=get_heading_style_name(document, level))
    remove_paragraph_numbering(paragraph)
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(6 if level == 1 else 3)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = True
    font_name = "楷体_GB2312" if level == 2 else "仿宋" if level == 3 else "黑体"
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 0, 0)


def chinese_section_number(index: int) -> str:
    numerals = "零一二三四五六七八九十"
    if index <= 10:
        return numerals[index]
    if index < 20:
        return f"十{numerals[index - 10]}"
    tens, ones = divmod(index, 10)
    suffix = numerals[ones] if ones else ""
    return f"{numerals[tens]}十{suffix}"


def is_unopened_record(record: DailyRecord, profile: dict) -> bool:
    category = profile["unopened_category"]
    return record.issue_counts.get(category, 0) > 0 and sum(record.issue_counts.values()) == 1


def record_status_text(record: DailyRecord) -> str:
    return normalize_issue_text(record.issue_text)


def is_closed_status_record(record: DailyRecord, profile: dict) -> bool:
    status_text = record_status_text(record)
    return any(keyword in status_text for keyword in profile.get("non_problem_status_keywords", ()))


def is_body_record(record: DailyRecord, profile: dict) -> bool:
    if not record.has_problem or is_unopened_record(record, profile):
        return False
    if sum(record.issue_counts.values()) == 0 and is_closed_status_record(record, profile):
        return False
    return True


def record_has_report_problem(record: DailyRecord, profile: dict | None = None) -> bool:
    if profile and is_closed_status_record(record, profile):
        return False
    return sum(record.issue_counts.values()) > 0


def format_station_issue_text(record: DailyRecord, profile: dict | None = None) -> tuple[str, bool]:
    if record_has_report_problem(record, profile):
        return format_issue_items(split_issue_items(record.issue_text)), True
    return normalize_issue_text(record.issue_text), False


def format_record_date_code(record: DailyRecord) -> str:
    return f"{record.month:02d}{record.day:02d}"


def format_station_label(record: DailyRecord, station_index: int, use_station_numbers: bool, station_type: str) -> str:
    if normalize_station_type(station_type) == "中转站":
        return f"{record.station}（{format_record_date_code(record)}）"
    return f"{station_index}.{record.station}" if use_station_numbers else record.station


def build_body_records(
    records: list[DailyRecord],
    station_type: str,
    categories: list[tuple[str, tuple[str, ...]]],
) -> list[DailyRecord]:
    if normalize_station_type(station_type) == "中转站":
        # 同一站点同一天只保留一条记录，避免日报中不同前缀格式造成重复
        seen: dict[tuple[str, str, int, int], DailyRecord] = {}
        for record in sorted(records, key=lambda r: (r.street, r.month, r.day, r.station)):
            key = (record.street, record.station, record.month, record.day)
            if key not in seen:
                seen[key] = record
        return list(seen.values())
    return select_latest_station_records(records, categories)


def limit_image_rows(image_rows: tuple[ImageRow, ...], max_images: int) -> tuple[ImageRow, ...]:
    images = [image for image_row in image_rows for image in image_row.images][:max_images]
    if not images:
        return ()
    return (ImageRow(images=tuple(images)),)


def normalize_display_issue_text(text: str) -> str:
    text = text.strip()
    if re.match(r"^（\d+）", text):
        return text
    return f"（1）{text}"


def split_issue_items(text: str) -> list[str]:
    normalized = normalize_issue_text(text)
    matches = list(re.finditer(r"（\d+）", normalized))
    if not matches:
        return [strip_trailing_punctuation(normalized)]
    items: list[str] = []
    for index, match in enumerate(matches):
        start = match.end()
        end = matches[index + 1].start() if index + 1 < len(matches) else len(normalized)
        item = strip_trailing_punctuation(normalized[start:end].strip())
        if item:
            items.append(item)
    return items


def strip_trailing_punctuation(text: str) -> str:
    return text.strip().rstrip("；;。")


def format_issue_items(items: list[str]) -> str:
    unique_items: list[str] = []
    seen: set[str] = set()
    for item in items:
        normalized = strip_trailing_punctuation(item)
        if normalized and normalized not in seen:
            unique_items.append(normalized)
            seen.add(normalized)
    return "；".join(f"（{index}）{item}" for index, item in enumerate(unique_items, start=1)) + "。"


def issue_item_key(item: str, categories: list[tuple[str, tuple[str, ...]]]) -> str:
    category = best_issue_category(item, categories)
    if category:
        return f"category:{category}"
    return f"text:{item}"


def image_row_signature(row: ImageRow) -> tuple[tuple[int, bytes], ...]:
    return tuple((len(image.blob), image.blob[:32]) for image in row.images)


def merge_street_records(records: list[DailyRecord], categories: list[tuple[str, tuple[str, ...]]]) -> list[MergedRecord]:
    merged_items: dict[str, list[str]] = {}
    merged_texts: dict[str, dict[str, str]] = {}
    merged_images: dict[str, dict[str, tuple[ImageRow, ...]]] = {}

    for record in records:
        issue_keys = merged_items.setdefault(record.station, [])
        issue_texts = merged_texts.setdefault(record.station, {})
        issue_images = merged_images.setdefault(record.station, {})

        for item in split_issue_items(record.issue_text):
            key = issue_item_key(item, categories)
            if key not in issue_texts:
                issue_keys.append(key)
                issue_texts[key] = item
            # Records are sorted by date, so assignment keeps only the latest
            # image rows for duplicated problem categories.
            issue_images[key] = record.image_rows

    result: list[MergedRecord] = []
    for station, issue_keys in merged_items.items():
        rows: list[ImageRow] = []
        seen_rows: set[tuple[tuple[int, bytes], ...]] = set()
        for key in issue_keys:
            for row in merged_images.get(station, {}).get(key, ()):
                signature = image_row_signature(row)
                if signature in seen_rows:
                    continue
                rows.append(row)
                seen_rows.add(signature)
        result.append(
            MergedRecord(
                station=station,
                issue_texts=[merged_texts[station][key] for key in issue_keys],
                image_rows=rows,
            )
        )
    return result


def select_latest_station_records(
    records: list[DailyRecord],
    categories: list[tuple[str, tuple[str, ...]]],
) -> list[DailyRecord]:
    profile = get_station_profile("清洁站")
    latest_by_station: dict[tuple[str, str], list[DailyRecord]] = defaultdict(list)
    for record in records:
        latest_by_station[(record.street, record.station)].append(record)

    selected: list[DailyRecord] = []
    for station_key, station_records in latest_by_station.items():
        station_records = sorted(station_records, key=lambda record: (record.month, record.day))
        latest_month = station_records[-1].month
        latest_day = station_records[-1].day
        latest_records = [
            record for record in station_records if record.month == latest_month and record.day == latest_day
        ]

        closed_record = next((record for record in reversed(latest_records) if is_closed_status_record(record, profile)), None)
        if closed_record is not None:
            continue

        problem_record = next(
            (record for record in reversed(latest_records) if record_has_report_problem(record, profile)),
            None,
        )
        if problem_record is not None:
            selected.append(problem_record)
            continue

        selected.append(latest_records[-1])

    return sorted(selected, key=lambda record: (record.street, record.station))


def summarize_records(
    records: list[DailyRecord],
    categories: list[tuple[str, tuple[str, ...]]],
) -> tuple[dict[str, list[DailyRecord]], dict[str, dict[str, int]]]:
    by_street: dict[str, list[DailyRecord]] = defaultdict(list)
    counts: dict[str, dict[str, int]] = defaultdict(lambda: {category: 0 for category, _ in categories})

    for record in records:
        street = normalize_street_name(record.street)
        by_street[street].append(record)
        for category, count in record.issue_counts.items():
            counts[street][category] += count

    return dict(sorted(by_street.items())), dict(sorted(counts.items()))


def build_summary_sentence(
    records: list[DailyRecord],
    year: int,
    month: int,
    profile: dict,
    street_count: int,
) -> str:
    start, end = report_period(year, month)
    problem_records = [record for record in records if record_has_report_problem(record, profile)]
    start_label = format_date_cn(start)
    end_label = format_date_cn(end)

    if start_label == end_label:
        prefix = f"{start_label}对本区{street_count}个街道{profile['subject']}进行了{len(records)}个次检查。"
    else:
        prefix = f"{start_label}至{end_label}对本区{street_count}个街道{profile['subject']}进行了{len(records)}个次检查。"

    if not problem_records:
        return f"{prefix}除未开门运行情况外，检查未发现其他问题。"
    return f"{prefix}检查发现需通报问题站次{len(problem_records)}个。"


def add_issue_table(document: Document, counts: dict[str, dict[str, int]], categories: list[tuple[str, tuple[str, ...]]]) -> None:
    headers = ["街道名称", *[category for category, _ in categories], "问题总数"]
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    set_table_row_min_height(table.rows[0])

    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_text(cell, header, bold=True)
        set_cell_shading(cell, "D9EAF7")

    total_by_category = {category: 0 for category, _ in categories}
    for street_counts in counts.values():
        for category, value in street_counts.items():
            total_by_category[category] += value

    rows = [("总计", total_by_category), *counts.items()]
    for street, street_counts in rows:
        row = table.add_row()
        set_table_row_min_height(row)
        values = [street]
        problem_total = 0
        for category, _ in categories:
            value = street_counts.get(category, 0)
            problem_total += value
            values.append("" if value == 0 else str(value))
        values.append("" if problem_total == 0 else str(problem_total))
        for index, value in enumerate(values):
            set_cell_text(row.cells[index], value)


def create_monthly_docx(
    records: list[DailyRecord],
    output_path: Path,
    year: int,
    month: int,
    station_type: str,
    template_path: Path | None = None,
    work_dir: Path | None = None,
) -> Path:
    profile = get_station_profile(station_type)
    categories = profile["categories"]
    if not records:
        raise ValueError(f"未找到 {month} 月{profile['subject']}日报记录")

    own_work_dir = work_dir is None
    if own_work_dir:
        work_dir = create_monthly_work_dir(output_path.parent, PROJECT_ROOT)

    try:
        document = create_document_from_template(template_path)
        jinja_path = export_jinja_text_template(template_path, station_type, profile, work_dir)
        if jinja_path:
            print(f"[临时] 已生成 Jinja2 文本模板：{jinja_path}")
        configure_styles(document)

        configure_portrait_a4_section(document.sections[0])

        context, image_map, counts = build_template_context(records, year, month, station_type, profile, categories)
        render_jinja_body(document, jinja_path, context, image_map)

        landscape = document.add_section(WD_ORIENT.LANDSCAPE)
        configure_landscape_a4_section(landscape)

        start, end = report_period(year, month)
        attachment = document.add_paragraph()
        attachment.alignment = WD_ALIGN_PARAGRAPH.CENTER
        attachment.paragraph_format.space_before = Pt(8)
        attachment.paragraph_format.space_after = Pt(8)
        run = attachment.add_run(
            f"附件：{format_date_cn(start)}至{format_date_cn(end)}"
            f"{profile['attachment_subject']}各项指标问题数"
        )
        run.font.name = "仿宋"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
        run.font.size = Pt(14)

        add_issue_table(document, counts, categories)

        cleanup_empty_paragraphs(document)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        try:
            document.save(output_path)
            return output_path
        except PermissionError:
            fallback_path = next_available_output_path(output_path)
            document.save(fallback_path)
            print(f"[提示] 输出文件可能已被打开或占用，已另存为：{fallback_path}")
            return fallback_path
    finally:
        if own_work_dir:
            cleanup_work_dir(work_dir)


def next_available_output_path(path: Path) -> Path:
    for index in range(1, 100):
        candidate = path.with_name(f"{path.stem}_{index}{path.suffix}")
        if not candidate.exists():
            return candidate
    raise PermissionError(f"输出文件被占用，且无法生成可用的新文件名：{path}")


def generate_monthly_report(
    raw_input_dir: Path = DEFAULT_RAW_INPUT_DIR,
    converted_dir: Path = DEFAULT_CONVERTED_DIR,
    output_dir: Path = DEFAULT_OUTPUT_DIR,
    month: str | int | None = None,
    year: int | None = None,
    station_type: str = DEFAULT_STATION_TYPE,
    engine: str = "auto",
    skip_convert: bool = False,
    template_path: Path | None = None,
    work_dir: Path | None = None,
) -> Path:
    own_work_dir = work_dir is None
    if own_work_dir:
        work_dir = create_monthly_work_dir(output_dir, converted_dir, PROJECT_ROOT)
    station_type = normalize_station_type(station_type)
    try:
        report_month = parse_month(month)
        report_year = parse_year(year, month)

        matching_doc_sources = [path for path in iter_doc_files(raw_input_dir) if path_matches_station_type(path, station_type)]

        if not skip_convert:
            summary = ensure_converted(raw_input_dir, converted_dir, station_type, engine, report_year, report_month)
            if summary.failed:
                messages = "\n".join(f"{path}: {error}" for path, error in summary.failed)
                raise ConversionError(f"部分日报转换失败：\n{messages}")

        record_sources = [raw_input_dir]
        if matching_doc_sources:
            record_sources.append(converted_dir / station_type)
        records = collect_records(record_sources, station_type, report_year, report_month)
        profile = get_station_profile(station_type)
        output_path = output_dir / profile["output_name"].format(year=report_year, month=report_month)
        return create_monthly_docx(records, output_path, report_year, report_month, station_type, template_path, work_dir)
    finally:
        if own_work_dir:
            cleanup_work_dir(work_dir)


def generate_monthly_reports(
    raw_input_dir: Path = DEFAULT_RAW_INPUT_DIR,
    converted_dir: Path = DEFAULT_CONVERTED_DIR,
    output_dir: Path = DEFAULT_OUTPUT_DIR,
    month: str | int | None = None,
    year: int | None = None,
    station_types: list[str] | tuple[str, ...] | None = None,
    engine: str = "auto",
    skip_convert: bool = False,
    template_paths: dict[str, Path | None] | None = None,
) -> list[Path]:
    selected_types = [normalize_station_type(value) for value in (station_types or [DEFAULT_STATION_TYPE])]
    selected_types = list(dict.fromkeys(selected_types))
    generated: list[Path] = []
    failures: list[str] = []
    work_dir = create_monthly_work_dir(output_dir, converted_dir, PROJECT_ROOT)

    try:
        for station_type in selected_types:
            template_path = (template_paths or {}).get(station_type)
            try:
                generated.append(
                    generate_monthly_report(
                        raw_input_dir=raw_input_dir,
                        converted_dir=converted_dir,
                        output_dir=output_dir,
                        month=month,
                        year=year,
                        station_type=station_type,
                        engine=engine,
                        skip_convert=skip_convert,
                        template_path=template_path,
                        work_dir=work_dir,
                    )
                )
            except ValueError as exc:
                failures.append(f"{station_type}: {exc}")

        if not generated:
            raise ValueError("未生成任何月报。" + ("\n" + "\n".join(failures) if failures else ""))
        if failures:
            print("[跳过]\n" + "\n".join(failures))
        return generated
    finally:
        cleanup_work_dir(work_dir)


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="汇总日报并生成密闭式清洁站月度检查通报")
    parser.add_argument("--raw-input", type=Path, default=DEFAULT_RAW_INPUT_DIR, help="原始 .doc 日报根目录")
    parser.add_argument("--converted-output", type=Path, default=DEFAULT_CONVERTED_DIR, help="转换后 .docx 日报根目录")
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT_DIR, help="月报输出目录")
    parser.add_argument("--month", required=True, help="月份，例如 4、04 或 2026-04")
    parser.add_argument("--year", type=int, help="报告年份，默认从 --month 解析，解析不到则为 2026")
    parser.add_argument("--station-type", nargs="+", default=[DEFAULT_STATION_TYPE], help="日报类型，可同时指定：清洁站 中转站")
    parser.add_argument("--template", type=Path, help="月报 docx 模板文件，可选")
    parser.add_argument(
        "--engine",
        choices=("auto", "word", "libreoffice"),
        default="auto",
        help="doc 转 docx 引擎，默认 auto",
    )
    parser.add_argument("--skip-convert", action="store_true", help="跳过 doc 转 docx 检查，直接汇总现有 docx")
    return parser


def main() -> int:
    args = build_parser().parse_args()
    try:
        output_paths = generate_monthly_reports(
            raw_input_dir=args.raw_input,
            converted_dir=args.converted_output,
            output_dir=args.output,
            month=args.month,
            year=args.year,
            station_types=args.station_type,
            engine=args.engine,
            skip_convert=args.skip_convert,
            template_paths={normalize_station_type(value): args.template for value in args.station_type} if args.template else None,
        )
    except PermissionError as exc:
        print(f"[失败] 无法写入目标文件，可能正在被 Word 打开：{exc.filename}", file=sys.stderr)
        print("[提示] 请关闭该文档后重新运行脚本，或通过 --output 指定其他输出目录。", file=sys.stderr)
        return 1
    except (ConversionError, FileNotFoundError, ValueError) as exc:
        print(f"[失败] {exc}", file=sys.stderr)
        return 1

    print("[完成] 月报已生成：")
    for output_path in output_paths:
        print(output_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
